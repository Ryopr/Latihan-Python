from docx import Document
from docxtpl import DocxTemplate

Alamat_Pengalaman="D:/Gdrive/2026/training/Latihan Python/Contoh data word/"

Nama_File_word="REF. 2019 MANUNGGAL DAYA KUKAR.docx"

def replace_text(doc, cari, ganti):
        # Ganti di paragraf biasa
        for para in doc.paragraphs:
                gabungan = para.text
                if cari in gabungan:
                    teks_baru = gabungan.replace(cari, ganti)
                    # Hapus semua runs
                    for i in range(len(para.runs)-1, -1, -1):
                        para.runs[i].clear()
                    # Tambahkan satu run baru
                    para.add_run(teks_baru)
# buka file docx tadi
document=Document(Alamat_Pengalaman+Nama_File_word)

# ganti file isi text di docx tadi
replace_text(document,"Priyo Hutomo ST.", "Jack, ST")


# save file docx yang dirubah tadi jadi file baru
document.save(Alamat_Pengalaman+"File Simpan Hasil edit.docx")